from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd
import json
from datetime import datetime

# ── Load data ──────────────────────────────────────────────
portfolio = pd.read_csv("portfolio_clean.csv")
payments  = pd.read_csv("payments.csv")
covenants = pd.read_csv("covenants.csv")

with open("portfolio_summary.json", "r") as f:
    summary = json.load(f)

with open("investor_report_may2026.txt", "r", encoding="utf-8") as f:
    investor_report = f.read()

with open("action_plan_may2026.txt", "r", encoding="utf-8") as f:
    action_plan = f.read()

# ── Create document ────────────────────────────────────────
doc = Document()

# ── Title ──────────────────────────────────────────────────
title = doc.add_heading('Impact Fund Portfolio Monitor', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Monthly Report — May 2026')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(13)
subtitle.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y")}').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

# ── Section 1: Portfolio Overview ─────────────────────────
doc.add_heading('1. Portfolio Overview', level=1)

table = doc.add_table(rows=7, cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

metrics = [
    ('Total Loans',          str(summary['total_loans'])),
    ('Total AUM',            f"${summary['total_aum_usd']:,.0f}"),
    ('Collection Rate',      f"{summary['collection_rate']}%"),
    ('On-Time Payment Rate', f"{summary['on_time_pct']}%"),
    ('Covenant Breaches',    str(summary['covenant_breaches'])),
    ('Missed Payments',      str(summary['missed_payments'])),
    ('Health Score',         f"{summary['health_score']} / 100 — WATCH ⚠"),
]

for i, (label, value) in enumerate(metrics):
    row = table.rows[i]
    row.cells[0].text = label
    row.cells[1].text = value
    row.cells[0].paragraphs[0].runs[0].font.bold = True

doc.add_paragraph('')

# ── Section 2: Charts ──────────────────────────────────────
doc.add_heading('2. Portfolio Analytics', level=1)

chart_files = [
    ("charts/health_metrics.png",   "Portfolio Health Metrics"),
    ("charts/payment_status.png",   "Payment Status Distribution"),
    ("charts/covenant_status.png",  "Covenant Compliance Status"),
    ("charts/loans_by_sector.png",  "Loans by Sector"),
    ("charts/loans_by_country.png", "Top 10 Countries by Loan Count"),
]

for chart_path, chart_title in chart_files:
    doc.add_paragraph(chart_title, style='Heading 3')
    doc.add_picture(chart_path, width=Inches(5.5))
    doc.add_paragraph('')

# ── Section 3: Investor Report ─────────────────────────────
doc.add_heading('3. Investor Report', level=1)
for line in investor_report.split('\n'):
    if line.strip():
        doc.add_paragraph(line.strip())

doc.add_paragraph('')

# ── Section 4: Action Plan ─────────────────────────────────
doc.add_heading('4. Fund Manager Action Plan', level=1)
for line in action_plan.split('\n'):
    if line.strip():
        doc.add_paragraph(line.strip())

doc.add_paragraph('')

# ── Section 5: Top Covenant Breaches Table ─────────────────
doc.add_heading('5. Covenant Breaches Detail', level=1)

breaches = covenants[covenants['covenant_status'] == 'Breach'][
    ['borrower_name', 'country', 'dscr', 'debt_to_equity', 'covenant_status']
].head(15)

table2 = doc.add_table(rows=len(breaches)+1, cols=5)
table2.style = 'Table Grid'

headers = ['Borrower', 'Country', 'DSCR', 'D/E Ratio', 'Status']
for i, h in enumerate(headers):
    cell = table2.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].font.bold = True

for row_idx, (_, row) in enumerate(breaches.iterrows(), start=1):
    table2.rows[row_idx].cells[0].text = str(row['borrower_name'])
    table2.rows[row_idx].cells[1].text = str(row['country'])
    table2.rows[row_idx].cells[2].text = str(row['dscr'])
    table2.rows[row_idx].cells[3].text = str(row['debt_to_equity'])
    table2.rows[row_idx].cells[4].text = str(row['covenant_status'])

doc.add_paragraph('')

# ── Section 6: Missed Payments Table ──────────────────────
doc.add_heading('6. Missed Payments Detail', level=1)

missed = payments[payments['payment_status'] == 'Missed'][
    ['borrower_name', 'country', 'due_date', 'amount_due']
].head(15)

table3 = doc.add_table(rows=len(missed)+1, cols=4)
table3.style = 'Table Grid'

headers3 = ['Borrower', 'Country', 'Due Date', 'Amount Due']
for i, h in enumerate(headers3):
    cell = table3.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].font.bold = True

for row_idx, (_, row) in enumerate(missed.iterrows(), start=1):
    table3.rows[row_idx].cells[0].text = str(row['borrower_name'])
    table3.rows[row_idx].cells[1].text = str(row['country'])
    table3.rows[row_idx].cells[2].text = str(row['due_date'])
    table3.rows[row_idx].cells[3].text = f"${float(row['amount_due']):,.2f}"

# ── Save document ──────────────────────────────────────────
output_file = "Impact_Fund_Portfolio_Report_May2026.docx"
doc.save(output_file)
print(f"Document saved: {output_file}")
print("Open it to see the full report with charts!")